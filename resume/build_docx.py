#!/usr/bin/env python3
"""resume.json -> resume/Shailesh_Chaudhari_Resume.docx (ATS-preferred format).

Run: <venv>/bin/python resume/build_docx.py   (needs python-docx)
"""
import json
from pathlib import Path

from docx import Document
from docx.shared import Pt

HERE = Path(__file__).parent
data = json.loads((HERE / "resume.json").read_text())

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)
for section in doc.sections:
    section.top_margin = section.bottom_margin = Pt(28)
    section.left_margin = section.right_margin = Pt(40)


def heading(text):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)


name = doc.add_paragraph()
r = name.add_run(data["name"])
r.bold = True
r.font.size = Pt(16)
doc.add_paragraph(data["title"])
c = data["contact"]
doc.add_paragraph(
    f"{c['email']} | {c['phone']} | {c['github']} | {c['linkedin']} | {c['portfolio']}"
)

heading("Summary")
doc.add_paragraph(data["summary"])

heading("Experience")
for e in data["experience"]:
    p = doc.add_paragraph()
    p.add_run(f"{e['title']}, {e['company']}").bold = True
    p.add_run(f"  ({e['start']} - {e['end']})")
    for b in e["bullets"]:
        doc.add_paragraph(b, style="List Bullet")

heading("Projects")
for pr in data["projects"]:
    p = doc.add_paragraph()
    p.add_run(pr["name"]).bold = True
    p.add_run(f" - {pr['tagline']} ({', '.join(pr['links'])})")
    for b in pr["bullets"]:
        doc.add_paragraph(b, style="List Bullet")

heading("Skills")
for s in data["skills"]:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(f"{s['label']}: ").bold = True
    p.add_run(s["items"])

heading("Education")
for e in data["education"]:
    doc.add_paragraph(e, style="List Bullet")

heading("Achievements")
for a in data["achievements"]:
    doc.add_paragraph(a, style="List Bullet")

out = HERE / "Shailesh_Chaudhari_Resume.docx"
doc.save(out)
print(f"built: {out}")
